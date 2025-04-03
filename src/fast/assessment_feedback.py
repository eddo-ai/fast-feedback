# assessment_feedback.py

import streamlit as st
import pandas as pd
from langchain import hub
from langchain_openai import ChatOpenAI
import docx
import io
import asyncio


st.title("Assessment Feedback")

# Initialize all session state variables
if "column_config" not in st.session_state:
    st.session_state.column_config = None
if "selected_columns" not in st.session_state:
    st.session_state.selected_columns = None
if "test_size" not in st.session_state:
    st.session_state.test_size = None
if "df" not in st.session_state:
    st.session_state.df = None
if "processed_df" not in st.session_state:
    st.session_state.processed_df = None
if "full_results_df" not in st.session_state:
    st.session_state.full_results_df = None
if "student_instructions" not in st.session_state:
    st.session_state.student_instructions = ""
if "grading_instructions" not in st.session_state:
    st.session_state.grading_instructions = ""
if "previous_selection" not in st.session_state:
    st.session_state.previous_selection = []


def prepare_batch(df, selected_columns, start_idx=0, batch_size=None):
    """Prepare a batch of samples for processing."""
    if not selected_columns:
        raise ValueError("No columns selected for student work")

    if not isinstance(df, pd.DataFrame):
        raise ValueError("Invalid dataframe provided")

    if batch_size is None:
        batch_size = len(df) - start_idx

    end_idx = min(start_idx + batch_size, len(df))
    batch_df = df.iloc[start_idx:end_idx].copy()

    # Validate selected columns exist
    missing_cols = [col for col in selected_columns if col not in batch_df.columns]
    if missing_cols:
        raise ValueError(f"Selected columns not found: {', '.join(missing_cols)}")

    student_instructions = st.session_state.get("student_instructions", "")
    grading_instructions = st.session_state.get("grading_instructions", "")

    # Prepare inputs for the batch with better error handling
    inputs = []
    valid_indices = []
    skipped_count = 0

    for idx, row in batch_df.iterrows():
        # Concatenate content from all selected columns
        student_work_parts = []
        for col in selected_columns:
            content = row[col]
            if not pd.isna(content) and str(content).strip():
                student_work_parts.append(f"{col}: {str(content).strip()}")

        # Skip if all parts are empty
        if not student_work_parts:
            skipped_count += 1
            continue

        # Join parts with newlines
        student_work = "\n\n".join(student_work_parts)

        inputs.append(
            {
                "student_work": student_work,
                "student_instructions": student_instructions,
                "grading_instructions": grading_instructions,
            }
        )
        valid_indices.append(idx)

    if skipped_count > 0:
        st.warning(f"Skipped {skipped_count} empty responses")

    if not inputs:
        raise ValueError("No valid student work found in selected columns")

    # Return only the rows with valid responses
    return batch_df.loc[valid_indices], inputs


async def process_batch(chain, batch_df, inputs):
    """Process a batch of samples and return results."""
    try:
        # Process the batch
        results = await chain.abatch(inputs)

        # Get all unique keys from the results excluding input fields
        result_keys = set()
        for result in results:
            result_keys.update(result.keys())
        result_keys = result_keys - {
            "student_instructions",
            "grading_instructions",
            "student_work",
        }

        # Dynamically add result columns to the dataframe
        for key in result_keys:
            batch_df[key] = [r.get(key, "") for r in results]

        return batch_df, results, True

    except Exception as e:
        st.error(f"Error processing batch: {str(e)}")
        return batch_df, None, False


def get_chain():
    """Initialize and return the LangChain chain."""
    if "chain" not in st.session_state:
        prompt = hub.pull("hey-aw/assessment_student_feedback")
        model = ChatOpenAI(model="gpt-4o")
        st.session_state.chain = prompt | model
    return st.session_state.chain


async def generate_feedback(df, selected_columns):
    """Main feedback generation function."""
    chain = get_chain()
    batch_df, inputs = prepare_batch(
        df, selected_columns, batch_size=st.session_state.test_size
    )
    processed_df, results, success = await process_batch(chain, batch_df, inputs)

    if success:
        return processed_df, results
    return None, None


async def process_remaining(df, selected_columns, test_size):
    """Process remaining samples after test batch."""
    chain = get_chain()

    # Validate inputs
    if not selected_columns:
        st.error("No column selected for student work")
        return None

    try:
        batch_df, inputs = prepare_batch(
            df,
            selected_columns,
            start_idx=test_size,
            batch_size=len(df) - test_size,  # Process all remaining
        )
        processed_df, results, success = await process_batch(chain, batch_df, inputs)

        if success:
            # Create a new dataframe with all columns
            full_df = df.copy()

            # Add new columns from results if they don't exist
            for col in processed_df.columns:
                if col not in full_df.columns:
                    full_df[col] = None

            # Update processed rows
            full_df.loc[processed_df.index] = processed_df
            return full_df

    except Exception as e:
        st.error(f"Error processing responses: {str(e)}")

    return None


def read_file_upload(uploaded_file):
    if uploaded_file is not None:
        # Get file extension
        file_ext = uploaded_file.name.split(".")[-1].lower()

        if file_ext == "docx":
            # Read DOCX files
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            content = []

            for element in doc.element.body:
                if element.tag.endswith("p"):  # Paragraph
                    paragraph = docx.Document(
                        io.BytesIO(uploaded_file.getvalue())
                    ).paragraphs[len(content)]
                    if paragraph.text.strip():
                        content.append(paragraph.text)

                elif element.tag.endswith("tbl"):  # Table
                    table = doc.tables[
                        len(
                            [
                                e
                                for e in doc.element.body[
                                    : doc.element.body.index(element)
                                ]
                                if e.tag.endswith("tbl")
                            ]
                        )
                    ]
                    table_data = []

                    # Get headers
                    headers = []
                    for cell in table.rows[0].cells:
                        headers.append(cell.text.strip())

                    # Get data rows
                    for row in table.rows[1:]:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        if any(cell for cell in row_data):  # Skip empty rows
                            table_data.append(dict(zip(headers, row_data)))

                    # Format table as text
                    if table_data:
                        content.append("\nTable:")
                        content.append(" | ".join(headers))
                        content.append(
                            "-"
                            * (sum(len(h) for h in headers) + 3 * (len(headers) - 1))
                        )
                        for row in table_data:
                            content.append(" | ".join(row.values()))
                        content.append("")

            return "\n".join(content)
        else:
            # Read TXT and MD files
            return uploaded_file.getvalue().decode("utf-8")
    return ""


async def handle_test_batch(df, selected_columns, test_size):
    """Handle the test batch processing."""
    test_df, results = await generate_feedback(df, selected_columns)

    if test_df is not None and results:
        # Create column config dynamically
        column_config = {}
        for col in test_df.columns:
            if col in selected_columns:
                continue
            if test_df[col].dtype == "object":
                column_config[col] = st.column_config.TextColumn(
                    col.replace("_", " ").title(), width="large"
                )
            else:
                column_config[col] = st.column_config.Column(
                    col.replace("_", " ").title()
                )

        # Store results in session state
        st.session_state.processed_df = test_df
        st.session_state.column_config = column_config
        return column_config
    return None


async def handle_remaining_samples(df, selected_columns, test_size, column_config):
    """Handle processing the remaining samples."""
    remaining_df = await process_remaining(df, selected_columns, test_size)
    if remaining_df is not None:
        st.write("### All Results")
        st.dataframe(remaining_df, column_config=column_config, hide_index=True)

        # Download all results
        csv_all = remaining_df.to_csv(index=False)
        st.download_button(
            label="Download All Results",
            data=csv_all,
            file_name="all_results.csv",
            mime="text/csv",
        )

        # Save full results to session state
        st.session_state.full_results_df = remaining_df


async def process_full_batch(df, selected_columns):
    """Process the full batch of responses."""
    # Process a small batch first to get the column structure
    test_df, results = await generate_feedback(df, selected_columns)

    if test_df is not None and results:
        # Create column config based on the test results
        column_config = {}
        for col in test_df.columns:
            if col in selected_columns:
                continue
            if test_df[col].dtype == "object":
                column_config[col] = st.column_config.TextColumn(
                    col.replace("_", " ").title(), width="large"
                )
            else:
                column_config[col] = st.column_config.Column(
                    col.replace("_", " ").title()
                )

        # Process all responses with the correct column config
        full_df = await process_remaining(
            df, selected_columns, 0
        )  # Start from beginning

        if full_df is not None:
            st.write("### All Results")
            st.dataframe(full_df, column_config=column_config, hide_index=True)

            # Download all results
            csv_all = full_df.to_csv(index=False)
            st.download_button(
                label="Download All Results",
                data=csv_all,
                file_name="all_results.csv",
                mime="text/csv",
            )

            # Save full results to session state
            st.session_state.full_results_df = full_df
            return True

    return False


def get_feedback_filename(original_filename, suffix=""):
    """Generate a feedback filename based on the original file."""
    # Remove .csv extension if present
    base_name = original_filename.rsplit(".", 1)[0]
    # Add feedback prefix and suffix
    return f"Feedback - {base_name}{suffix}.csv"


# Instructions setup section
st.header("1. Criteria")


st.write("### Student Instructions")
instructions_file = st.file_uploader(
    "Upload Student Instructions",
    type=["txt", "md", "docx"],
)
if instructions_file:
    st.session_state.student_instructions = read_file_upload(instructions_file)
    with st.expander("Instructions preview"):
        st.write(st.session_state.student_instructions)

st.write("### Grading Rubric")
rubric_file = st.file_uploader(
    "Upload Grading Instructions and Rubric",
    type=["txt", "md", "docx"],
)
if rubric_file:
    st.session_state.grading_instructions = read_file_upload(rubric_file)
    with st.expander("Rubric preview"):
        st.write(st.session_state.grading_instructions)

# Process responses section
if not (
    st.session_state.get("student_instructions")
    or st.session_state.get("grading_instructions")
):
    st.stop()
else:
    st.header("2. Process Student Responses")
    responses_file = st.file_uploader("Upload Student Responses (CSV)", type=["csv"])

    if responses_file is not None:
        df = pd.read_csv(responses_file)
        st.write(f"Loaded {len(df)} responses")

        # Column selection first
        st.write("### Select Response Columns")

        # Get previous selection if it exists and is valid
        previous_selection = st.session_state.previous_selection
        default_selection = [col for col in previous_selection if col in df.columns]

        selected_columns = st.multiselect(
            "Select columns containing student responses",
            df.columns,
            default=default_selection,
            help="Choose one or more columns that contain student work to be graded. Multiple columns will be combined.",
        )

        if selected_columns:
            st.session_state.selected_columns = selected_columns

            # Show data preview of selected columns
            st.write("### Student Responses Preview")
            st.dataframe(df[selected_columns], hide_index=True)

            # Step 3: Process Samples
            st.subheader("3. Process Samples")

            # Show processing options
            col1, col2 = st.columns(2)

            with col1:
                with st.form("test_sample"):
                    st.write("Test with a sample batch")
                    test_size = st.number_input(
                        "Sample size",
                        min_value=1,
                        max_value=len(df),
                        value=min(5, len(df)),
                    )
                    submitted = st.form_submit_button("Run Test Sample", type="primary")

            with col2:
                with st.form("full_batch"):
                    st.write("Process all responses")
                    st.write(f"Total responses: {len(df)}")
                    process_all = st.form_submit_button("Process All", type="primary")

            # Handle test sample processing
            if submitted:
                with st.spinner(
                    f"Processing {test_size} samples from column: {selected_columns[0]}..."
                ):
                    try:
                        column_config = asyncio.run(
                            handle_test_batch(
                                df, st.session_state.selected_columns, test_size
                            )
                        )

                        if column_config is not None:
                            st.write("### Results")
                            st.dataframe(
                                st.session_state.processed_df,
                                column_config=column_config,
                                hide_index=True,
                            )

                            csv_test = st.session_state.processed_df.to_csv(index=False)
                            feedback_filename = get_feedback_filename(
                                responses_file.name, f" - {test_size} samples"
                            )
                            st.download_button(
                                label="Download Results",
                                data=csv_test,
                                file_name=feedback_filename,
                                mime="text/csv",
                            )
                        else:
                            st.error("Failed to process test batch")
                    except Exception as e:
                        st.error(f"Error processing test batch: {str(e)}")

            # Handle full batch processing
            if process_all:
                with st.spinner(
                    f"Processing {len(df)} responses from column: {selected_columns[0]}..."
                ):
                    try:
                        success = asyncio.run(
                            process_full_batch(df, st.session_state.selected_columns)
                        )
                        if success:
                            feedback_filename = get_feedback_filename(
                                responses_file.name
                            )
                            csv_all = st.session_state.full_results_df.to_csv(
                                index=False
                            )
                            st.download_button(
                                label="Download All Results",
                                data=csv_all,
                                file_name=feedback_filename,
                                mime="text/csv",
                            )
                            st.success("Processing complete!")
                    except Exception as e:
                        st.error(f"Error processing responses: {str(e)}")
